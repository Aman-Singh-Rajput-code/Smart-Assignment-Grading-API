import PyPDF2
import docx
import re
import logging
import traceback

logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_path):
    """Extract text content from a PDF file."""
    text_list = []
    try:
        logger.info(f"Extracting text from PDF: {pdf_path}")
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            logger.info(f"PDF has {len(pdf_reader.pages)} pages")
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_list.append(page_text)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        logger.error(traceback.format_exc())

    text = "\n".join(text_list)
    logger.info(f"Extracted Text Length: {len(text)}")
    logger.debug(f"Text Preview: {text[:1000]}")
    return text  # ✅ returns a single string


def extract_text_from_docx(docx_path):
    """Extract text content from a DOCX file."""
    text = ""
    try:
        logger.info(f"Extracting text from DOCX: {docx_path}")
        doc = docx.Document(docx_path)
        logger.info(f"DOCX has {len(doc.paragraphs)} paragraphs")
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        logger.error(traceback.format_exc())

    logger.info(f"Total extracted text length: {len(text)} characters")
    logger.debug(f"First 200 characters of extracted text: {text[:200]}")
    return text  # ✅ ensures a single string is returned


def extract_qa_pairs(text):
    """
    Extract question-answer pairs from document text.

    Returns a list of dictionaries with question and answer.
    """
    logger.info("Extracting QA pairs from document text")
    qa_pairs = []

    logger.debug(f"First 500 characters of text: {text[:500]}")

    # Pattern 1: Q1: ... A1: ...
    pattern1 = r'(?:Q|Question)\s*(\d+)[\s:\.]*(.*?)(?:\n|\r\n?)(?:A|Answer)\s*\1[\s:\.]*(.*?)(?=(?:Q|Question)\s*\d+|\Z)'
    logger.debug(f"Trying pattern 1: {pattern1}")

    pattern2 = r'(\d+)[\.:\)]\s*(.*?)(?:\n|\r\n?)(?:A|Answer|Ans)[\.:\s]*(.*?)(?=\d+[\.:\)]\s*|\Z)'
    logger.debug(f"Trying pattern 2: {pattern2}")

    matches = re.findall(pattern1, text, re.DOTALL | re.IGNORECASE)
    logger.debug(f"Pattern 1 found {len(matches)} matches")

    if matches:
        for match in matches:
            qa_pair = {
                'question_num': match[0],
                'question': match[1].strip(),
                'answer': match[2].strip()
            }
            qa_pairs.append(qa_pair)
            logger.debug(f"Found QA pair: Q{match[0]}: {match[1][:30]}... A: {match[2][:30]}...")
    else:
        matches = re.findall(pattern2, text, re.DOTALL | re.IGNORECASE)
        logger.debug(f"Pattern 2 found {len(matches)} matches")

        if matches:
            for match in matches:
                qa_pair = {
                    'question_num': match[0],
                    'question': match[1].strip(),
                    'answer': match[2].strip()
                }
                qa_pairs.append(qa_pair)
                logger.debug(f"Found QA pair: Q{match[0]}: {match[1][:30]}... A: {match[2][:30]}...")

    # Fallback line-by-line enhanced extraction
    if not qa_pairs:
        logger.info("No structured QA pairs found with regex, trying enhanced line-by-line approach")
        lines = text.split('\n')
        current_question = None
        current_q_num = None

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            if i < 10:
                logger.debug(f"Line {i}: {line}")

            q_match = re.match(r'(?:Q|Question)?\s*(\d+)[\.:)\s]+(.+)', line, re.IGNORECASE)
            if q_match:
                current_q_num = q_match.group(1)
                current_question = q_match.group(2).strip()
                continue

            a_match = re.match(r'(?:A|Answer|Ans)[\.:)\s]+(.+)', line, re.IGNORECASE)
            if current_question and a_match:
                answer = a_match.group(1).strip()
                qa_pairs.append({
                    'question_num': current_q_num,
                    'question': current_question,
                    'answer': answer
                })
                logger.debug(f"Found answer for Q{current_q_num}: {answer[:30]}...")
                current_question = None
                current_q_num = None

    logger.info(f"Extracted {len(qa_pairs)} QA pairs in total")
    return qa_pairs
